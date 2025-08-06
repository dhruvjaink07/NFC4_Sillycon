# utils/agents.py
import os
import json
import re
import time
from typing import List
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import fitz  # PyMuPDF
from langchain_google_genai import ChatGoogleGenerativeAI
from gliner import GLiNER
from dotenv import load_dotenv

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

try:
    GLiNER.from_pretrained("urchade/gliner_base")
except:
    gliner = None

try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0,
        google_api_key=GEMINI_API_KEY,
        timeout=30
    )
except:
    llm = None

class RunnerAgent:
    def load_text(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                self.json_data = json.load(f)
                text = json.dumps(self.json_data, indent=2)
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type.")
        return text

    def get_file_type(self, file_path: str) -> str:
        return file_path.split(".")[-1]

    def save_redacted_text(self, redacted_text: str, original_path: str, output_path: str):
        file_type = self.get_file_type(original_path)
        if file_type == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(redacted_text)
        elif file_type == "json":
            try:
                redacted_json = json.loads(redacted_text)
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(redacted_json, f, indent=2)
            except json.JSONDecodeError:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(redacted_text)
        elif file_type == "docx":
            doc = Document()
            for paragraph in redacted_text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(output_path)
        elif file_type == "pdf":
            c = canvas.Canvas(output_path, pagesize=letter)
            lines = redacted_text.split("\n")
            width, height = letter
            y = height - 40
            for line in lines:
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:100])
                y -= 15
            c.save()

class RedactorAgent:
    def __init__(self):
        self.name_exclusions = {"united states", "new york", "monday", "tuesday"}

    def detect_sensitive_info(self, text: str) -> List[dict]:
        results = []
        if gliner is not None:
            try:
                entities = gliner.predict_entities(text[:5000], labels=[
                    "Person", "Organization", "Date", "Email", "Phone",
                    "Location", "URL", "Money", "Time"
                ], threshold=0.4)
                for ent in entities:
                    label = ent["label"].lower()
                    val = ent["text"].strip()
                    if label == "person" and self._is_likely_person_name(val):
                        results.append({"type": "name", "value": val})
                    elif label in ["organization", "email", "phone", "location", "url", "date", "time", "money"]:
                        results.append({"type": label, "value": val})
            except:
                pass
        results += self._regex_fallback(text)
        unique = {f"{r['type']}_{r['value'].lower()}": r for r in results}
        return list(unique.values())

    def _regex_fallback(self, text: str) -> List[dict]:
        patterns = {
            "email": r"[\w\.-]+@[\w\.-]+",
            "phone": r"\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b",
            "url": r"https?://[^\s]+",
            "ssn": r"\b\d{3}-\d{2}-\d{4}\b",
            "credit_card": r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b"
        }
        found = []
        for typ, pat in patterns.items():
            matches = re.findall(pat, text)
            for m in matches:
                found.append({"type": typ, "value": m})
        return found

    def _is_likely_person_name(self, name: str) -> bool:
        return name.lower() not in self.name_exclusions and len(name.split()) in [2, 3]

    def redact(self, text: str, sensitive_items: List[dict]) -> str:
        sorted_items = sorted(sensitive_items, key=lambda x: len(x["value"]), reverse=True)
        for item in sorted_items:
            val = re.escape(item["value"])
            tag = f"[REDACTED_{item['type'].upper()}]"
            text = re.sub(val, tag, text, flags=re.IGNORECASE)
        return text

    def redact_pdf_pymupdf(self, file_path: str, sensitive_items: List[dict], output_path: str):
        try:
            doc = fitz.open(file_path)
            for page in doc:
                for item in sensitive_items:
                    areas = page.search_for(item["value"])
                    for inst in areas:
                        page.add_redact_annot(inst, fill=(0, 0, 0))
                page.apply_redactions()
            doc.save(output_path)
            doc.close()
        except:
            pass

class ComplianceAgent:
    def apply_policy(self, pii_items: list, compliance_type: str) -> list:
        redactions = []
        for item in pii_items:
            if compliance_type == "GDPR":
                redactions.append(item["value"])
            elif compliance_type == "HIPAA":
                if item["type"] in ["ssn", "phone", "name", "email"]:
                    redactions.append(item["value"])
            elif compliance_type == "DPDP":
                if item["type"] != "url":
                    redactions.append(item["value"])
        return redactions

    def validate_redaction(self, redacted_text: str, compliance_type: str) -> str:
        if llm is None:
            return f"Compliance validation skipped."
        try:
            prompt = (
                f"You are validating redaction under {compliance_type}.\n"
                f"Redacted Text: {redacted_text[:1000]}"
            )
            result = llm.invoke(prompt)
            return result.content if hasattr(result, 'content') else str(result)
        except:
            return "Compliance validation failed."

class AuditAgent:
    def log_metadata(self, original_text: str, redacted_text: str, sensitive_items: List[dict], compliance_feedback: str, file_path: str):
        metadata = {
            "timestamp": time.strftime('%Y-%m-%d %H:%M:%S'),
            "file_path": file_path,
            "original_length": len(original_text),
            "redacted_length": len(redacted_text),
            "redacted_count": len(sensitive_items),
            "compliance_notes": compliance_feedback
        }
        os.makedirs("audit_logs", exist_ok=True)
        log_file = f"audit_logs/audit_log_{os.path.basename(file_path)}_{int(time.time())}.json"
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(metadata, f, indent=2)

class CoordinatorAgent:
    def __init__(self):
        self.runner = RunnerAgent()
        self.redactor = RedactorAgent()
        self.compliance = ComplianceAgent()
        self.audit = AuditAgent()

    def handle_files(self, file_paths: List[str], compliance_type: str):
        for path in file_paths:
            try:
                text = self.runner.load_text(path)
                pii_items = self.redactor.detect_sensitive_info(text)
                redacted = self.redactor.redact(text, pii_items)
                ext = os.path.splitext(path)[1]
                output = path.replace(ext, f"_redacted{ext}")
                if ext == ".pdf":
                    self.redactor.redact_pdf_pymupdf(path, pii_items, output)
                else:
                    self.runner.save_redacted_text(redacted, path, output)
                feedback = self.compliance.validate_redaction(redacted, compliance_type)
                self.audit.log_metadata(text, redacted, pii_items, feedback, path)
                print(f"✅ Should have saved redacted file at: {output}")
                print(f"Exists? {os.path.exists(output)}")

            except:
                continue
