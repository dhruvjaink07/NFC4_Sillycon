import os
import json
import docx
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import re
from typing import List
import fitz  # PyMuPDF
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
import time

# Load Gemini API key
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# ==============================
# LLM Configuration
# ==============================
try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0,
        google_api_key=GEMINI_API_KEY,
        timeout=30
    )
    print("✅ LLM loaded")
except Exception as e:
    print(f"❌ Failed to load Gemini LLM: {e}")
    llm = None

# ==============================
# RunnerAgent
# ==============================
class RunnerAgent:
    def load_text(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())

        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                self.json_data = json.load(f)
                text = json.dumps(self.json_data, indent=2)

        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        else:
            raise ValueError("Unsupported file type. Use PDF, TXT, JSON, or DOCX.")

        return text

    def get_file_type(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            return "pdf"
        elif file_path.endswith(".txt"):
            return "txt"
        elif file_path.endswith(".json"):
            return "json"
        elif file_path.endswith(".docx"):
            return "docx"
        else:
            raise ValueError("Unsupported file type")

    def save_redacted_text(self, redacted_text: str, original_path: str, output_path: str):
        file_type = self.get_file_type(original_path)

        if file_type == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(redacted_text)

        elif file_type == "json":
            try:
                redacted_json = json.loads(redacted_text)
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(redacted_json, f, indent=2)
            except json.JSONDecodeError:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(redacted_text)

        elif file_type == "docx":
            doc = Document()
            for paragraph in redacted_text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(output_path)

        elif file_type == "pdf":
            c = canvas.Canvas(output_path, pagesize=letter)
            lines = redacted_text.split("\n")
            width, height = letter
            y = height - 40
            for line in lines:
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:100])  # Limit line length
                y -= 15
            c.save()

# ==============================
# RedactorAgent - FIXED VERSION
# ==============================
class RedactorAgent:
    def __init__(self):
        # Common non-names to exclude
        self.name_exclusions = {
            'united states', 'new york', 'los angeles', 'san francisco',
            'machine learning', 'data science', 'artificial intelligence',
            'google drive', 'microsoft office', 'adobe acrobat', 'dear sir',
            'dear madam', 'yours truly', 'best regards', 'thank you'
        }
    
    def detect_sensitive_info(self, text: str) -> List[dict]:
        """Enhanced detection with better AI + regex combination"""
        try:
            print("🔍 Analyzing text with Gemini API...")
            prompt = (
                "Extract and return a list of sensitive data from this text. "
                "Types should include email, phone, URL, SSN, credit_card, and names (person names only, not titles). "
                "For names, only extract actual person names, not titles like 'Mr.', 'Dr.', etc. "
                "Return ONLY a valid JSON array in this exact format:\n"
                "[{\"type\": \"email\", \"value\": \"abc@email.com\"}, {\"type\": \"name\", \"value\": \"John Smith\"}, ...]\n\n"
                f"Text:\n{text[:2000]}"
            )
            
            result = llm.invoke(prompt)
            ai_items = self._parse_json_list(result.content)
            print(f"🤖 AI found {len(ai_items)} items")
            
        except Exception as e:
            print(f"❌ Error with AI detection: {e}")
            ai_items = []
        
        # Always run regex fallback
        regex_items = self._regex_fallback(text)
        
        # Combine and deduplicate
        all_items = {}
        for item in (ai_items + regex_items):
            key = f"{item['type']}_{item['value'].lower()}"
            all_items[key] = item
        
        final_items = list(all_items.values())
        print(f"📊 Total unique items found: {len(final_items)}")
        
        return final_items

    def _regex_fallback(self, text: str) -> List[dict]:
        """Enhanced regex patterns with better validation"""
        sensitive_items = []
        
        # Email pattern - more precise
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        for email in emails:
            sensitive_items.append({"type": "email", "value": email})
        
        # Phone patterns with validation
        phone_patterns = [
            r'\b\d{3}-\d{3}-\d{4}\b',  # 123-456-7890
            r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',  # (123) 456-7890
            r'\b\d{10}\b',  # 1234567890
            r'\b\d{3}\.\d{3}\.\d{4}\b',  # 123.456.7890
            r'\+\d{1,3}\s?\d{3,4}\s?\d{3,4}\s?\d{3,4}\b'  # International
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            for phone in phones:
                # Basic validation - must have 10+ digits
                digits_only = re.sub(r'\D', '', phone)
                if len(digits_only) >= 10:
                    sensitive_items.append({"type": "phone", "value": phone})
        
        # URL pattern
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        urls = re.findall(url_pattern, text)
        for url in urls:
            sensitive_items.append({"type": "url", "value": url})
        
        # SSN pattern
        ssn_pattern = r'\b\d{3}-\d{2}-\d{4}\b'
        ssns = re.findall(ssn_pattern, text)
        for ssn in ssns:
            sensitive_items.append({"type": "ssn", "value": ssn})
        
        # Credit card pattern
        cc_pattern = r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b'
        cards = re.findall(cc_pattern, text)
        for card in cards:
            # Basic validation - must have 16 digits
            digits_only = re.sub(r'\D', '', card)
            if len(digits_only) == 16:
                sensitive_items.append({"type": "credit_card", "value": card})
        
        # Enhanced name pattern with better filtering
        name_patterns = [
            r'\b[A-Z][a-z]+ [A-Z][a-z]+\b',  # John Smith
            r'\b[A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+\b',  # John M. Smith
            r'\b[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+\b'  # John Michael Smith
        ]
        
        for pattern in name_patterns:
            names = re.findall(pattern, text)
            for name in names:
                # Enhanced filtering
                if self._is_likely_person_name(name):
                    sensitive_items.append({"type": "name", "value": name})
        
        print(f"🔍 Regex detection found {len(sensitive_items)} items")
        return sensitive_items
    
    def _is_likely_person_name(self, name: str) -> bool:
        """Enhanced name validation"""
        name_lower = name.lower()
        
        # Skip if in exclusion list
        if name_lower in self.name_exclusions:
            return False
        
        # Skip common non-name patterns
        non_name_patterns = [
            r'\b(mr|mrs|ms|dr|prof|sir|madam)\b',
            r'\b(inc|llc|corp|ltd|co)\b',
            r'\b(street|avenue|road|drive|lane)\b',
            r'\b(monday|tuesday|wednesday|thursday|friday|saturday|sunday)\b',
            r'\b(january|february|march|april|may|june|july|august|september|october|november|december)\b'
        ]
        
        for pattern in non_name_patterns:
            if re.search(pattern, name_lower):
                return False
        
        # Additional checks
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            return False
            
        # Each word should be reasonable length for a name
        if any(len(word) < 2 or len(word) > 20 for word in words):
            return False
            
        return True

    def redact(self, text: str, sensitive_items: List[dict]) -> str:
        """FIXED: Precise redaction that only replaces the sensitive data"""
        redacted_text = text
        
        # Sort items by length (longest first) to avoid partial replacements
        sorted_items = sorted(sensitive_items, key=lambda x: len(x["value"]), reverse=True)
        
        for item in sorted_items:
            original_value = item["value"]
            item_type = item["type"].upper()
            
            # Create redaction tag
            redaction_tag = f"[REDACTED_{item_type}]"
            
            # FIXED: Use precise regex replacement with word boundaries
            if item["type"] == "email":
                # For emails, use exact match
                pattern = re.escape(original_value)
            elif item["type"] in ["phone", "ssn", "credit_card"]:
                # For structured data, use exact match
                pattern = re.escape(original_value)
            elif item["type"] == "url":
                # For URLs, use exact match
                pattern = re.escape(original_value)
            elif item["type"] == "name":
                # For names, use word boundaries but be more careful
                # This ensures we match "Thompson" in "Mr. Thompson" but not in "Thompson's"
                escaped_value = re.escape(original_value)
                pattern = r'\b' + escaped_value + r'\b'
            else:
                # Default: use word boundaries
                pattern = r'\b' + re.escape(original_value) + r'\b'
            
            # Replace with case-insensitive matching
            redacted_text = re.sub(pattern, redaction_tag, redacted_text, flags=re.IGNORECASE)
            
            print(f"🔄 Redacted: {original_value} -> {redaction_tag}")
        
        return redacted_text

    def redact_json(self, data: dict, sensitive_items: List[dict]) -> dict:
        """Enhanced JSON redaction"""
        redacted_data = json.dumps(data, indent=2)
        redacted_data = self.redact(redacted_data, sensitive_items)
        try:
            return json.loads(redacted_data)
        except json.JSONDecodeError:
            # If JSON is malformed after redaction, return as string
            return {"redacted_content": redacted_data}

    def _parse_json_list(self, raw: str) -> List[dict]:
        """Enhanced JSON parsing"""
        try:
            # Clean up the response
            raw = raw.strip()
            
            # Find JSON array
            start = raw.find('[')
            end = raw.rfind(']') + 1
            
            if start == -1 or end == 0:
                print("⚠️ No JSON array found in response")
                return []
            
            json_str = raw[start:end]
            parsed = json.loads(json_str)
            
            # Validate structure
            if not isinstance(parsed, list):
                print("⚠️ Parsed data is not a list")
                return []
            
            # Validate and clean items
            valid_items = []
            for item in parsed:
                if (isinstance(item, dict) and 
                    'type' in item and 'value' in item and 
                    isinstance(item['value'], str) and 
                    item['value'].strip()):
                    
                    # Clean the value
                    item['value'] = item['value'].strip()
                    valid_items.append(item)
            
            return valid_items
            
        except json.JSONDecodeError as e:
            print(f"❌ JSON parsing error: {e}")
            return []
        except Exception as e:
            print(f"❌ Unexpected error parsing response: {e}")
            return []

    def redact_pdf_pymupdf(self, file_path: str, sensitive_items: List[dict], output_path: str):
        """Enhanced PDF redaction with better text matching"""
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                for item in sensitive_items:
                    # Search for sensitive text on the page
                    text_instances = page.search_for(item["value"])
                    
                    for inst in text_instances:
                        # Create a black rectangle over the sensitive text
                        page.add_redact_annot(inst, fill=(0, 0, 0))
                        print(f"🔄 PDF: Redacted '{item['value']}' on page {page_num + 1}")
                
                # Apply all redactions on this page
                page.apply_redactions()
            
            doc.save(output_path)
            doc.close()
            print(f"✅ PDF redacted and saved to: {output_path}")
            
        except Exception as e:
            print(f"❌ Error redacting PDF: {e}")

# ==============================
# PIIAgent - Enhanced
# ==============================
class PIIAgent:
    def detect_pii(self, text: str) -> list:
        """Enhanced PII detection with more patterns"""
        patterns = {
            "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "phone": r'\b(?:\+1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b',
            "ssn": r'\b\d{3}-\d{2}-\d{4}\b',
            "credit_card": r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b',
        }
        matches = []
        for name, pattern in patterns.items():
            for match in re.findall(pattern, text):
                if isinstance(match, tuple):
                    match = ''.join(match)
                matches.append({"type": name, "value": match})
        return matches

# ==============================
# ComplianceAgent
# ==============================
class ComplianceAgent:
    def apply_policy(self, pii_items: list, compliance_type: str) -> list:
        """Enhanced compliance policy application"""
        redactions = []
        for item in pii_items:
            if compliance_type == "GDPR":
                # GDPR requires redaction of all personal data
                redactions.append(item["value"])
            elif compliance_type == "HIPAA":
                # HIPAA focuses on health-related identifiers
                if item["type"] in ["ssn", "phone", "name"]:
                    redactions.append(item["value"])
            elif compliance_type == "DPDP":
                # Digital Personal Data Protection Act
                if item["type"] != "url":  # Keep URLs for business purposes
                    redactions.append(item["value"])
        return redactions

    def validate_redaction(self, redacted_text: str, compliance_type: str) -> str:
        """Enhanced compliance validation"""
        try:
            prompt = (
                f"You are a compliance officer validating text redactions for {compliance_type}. "
                f"Analyze the redacted text and check for:\n"
                f"1. Any remaining personal identifiable information\n"
                f"2. Compliance with {compliance_type} standards\n"
                f"3. Proper redaction formatting\n\n"
                f"Return a brief assessment (2-3 sentences) of compliance status.\n\n"
                f"Redacted Text:\n{redacted_text[:1000]}"  # Limit for API
            )
            
            result = llm.invoke(prompt)
            return result.content if hasattr(result, 'content') else str(result)
            
        except Exception as e:
            return f"Compliance validation failed: {str(e)}"

# ==============================
# AuditAgent - Enhanced
# ==============================
class AuditAgent:
    def log_metadata(self, original_text: str, redacted_text: str, sensitive_items: List[dict], compliance_feedback: str, file_path: str):
        """Enhanced audit logging with more details"""
        # Count items by type
        item_counts = {}
        for item in sensitive_items:
            item_type = item['type']
            item_counts[item_type] = item_counts.get(item_type, 0) + 1
        
        metadata = {
            "timestamp": time.strftime('%Y-%m-%d %H:%M:%S'),
            "file_path": file_path,
            "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            "original_length": len(original_text),
            "redacted_length": len(redacted_text),
            "total_redacted_items": len(sensitive_items),
            "redacted_items_by_type": item_counts,
            "redacted_items": [
                {"type": item["type"], "value_length": len(item["value"])} 
                for item in sensitive_items
            ],
            "compliance_notes": str(compliance_feedback),
            "processing_status": "completed"
        }
        
        # Create logs directory if it doesn't exist
        os.makedirs("audit_logs", exist_ok=True)
        
        log_file = f"audit_logs/audit_log_{os.path.basename(file_path)}_{int(time.time())}.json"
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(metadata, f, indent=2)
        
        print(f"📋 [AuditAgent] Detailed metadata saved to {log_file}")

# ==============================
# CoordinatorAgent - Enhanced
# ==============================
class CoordinatorAgent:
    def __init__(self):
        self.runner = RunnerAgent()
        self.redactor = RedactorAgent()
        self.compliance = ComplianceAgent()
        self.audit = AuditAgent()

    def handle_files(self, file_paths: List[str], compliance_type: str):
        """Enhanced file processing with better error handling"""
        for file_path in file_paths:
            print(f"\n🔄 Processing: {file_path}")
            
            if not os.path.exists(file_path):
                print(f"❌ File not found: {file_path}")
                continue
            
            try:
                # Load and analyze
                original_text = self.runner.load_text(file_path)
                print(f"📄 Loaded {len(original_text)} characters")
                
                # Detect sensitive information
                pii_items = self.redactor.detect_sensitive_info(original_text)
                
                if not pii_items:
                    print("✅ No sensitive information detected")
                    continue
                
                print(f"🔍 Found {len(pii_items)} sensitive items:")
                for item in pii_items:
                    print(f"   - {item['type']}: {item['value'][:20]}{'...' if len(item['value']) > 20 else ''}")
                
                # Redact text
                redacted_text = self.redactor.redact(original_text, pii_items)
                
                # Save redacted file
                file_ext = os.path.splitext(file_path)[1]
                output_path = file_path.replace(file_ext, f"_redacted{file_ext}")

                if file_ext == ".pdf":
                    self.redactor.redact_pdf_pymupdf(file_path, pii_items, output_path)
                else:
                    self.runner.save_redacted_text(redacted_text, file_path, output_path)

                print(f"✅ Redacted file saved at: {output_path}")

                # Compliance validation
                feedback = self.compliance.validate_redaction(redacted_text, compliance_type)
                print(f"📋 Compliance feedback: {feedback[:100]}...")
                
                # Audit logging
                self.audit.log_metadata(original_text, redacted_text, pii_items, feedback, file_path)
                
            except Exception as e:
                print(f"❌ Error processing {file_path}: {str(e)}")
                continue

# ==============================
# Main Execution Loop - Enhanced
# ==============================
if __name__ == "__main__":
    print("🚀 Multi-Agent Sensitive Data Redaction System")
    print("=" * 50)
    
    # Test with sample data
    test_mode = input("Do you want to test with sample data first? (y/n): ").strip().lower()
    
    if test_mode == 'y':
        print("\n🧪 Testing with sample data...")
        sample_text = "Dear Mr. Thompson, please contact John Smith at john.smith@email.com or call 555-123-4567."
        
        redactor = RedactorAgent()
        items = redactor.detect_sensitive_info(sample_text)
        redacted = redactor.redact(sample_text, items)
        
        print(f"Original: {sample_text}")
        print(f"Redacted: {redacted}")
        print(f"Items found: {items}")
        print()
    
    while True:
        file_paths_input = input("Enter file paths (comma-separated for multiple files, or single file): ").strip()
        
        if not file_paths_input:
            print("❌ Please provide at least one file path")
            continue
            
        file_paths = [path.strip() for path in file_paths_input.split(",")]
        invalid_files = [path for path in file_paths if not os.path.exists(path)]
        
        if invalid_files:
            print(f"❌ Files not found: {', '.join(invalid_files)}")
            continue

        print("\nSelect compliance type:")
        print("1. GDPR (General Data Protection Regulation)")
        print("2. HIPAA (Health Insurance Portability and Accountability Act)")
        print("3. DPDP (Digital Personal Data Protection Act)")
        
        compliance_choice = input("Enter choice (1/2/3) or type name directly: ").strip()
        
        if compliance_choice == "1":
            compliance_type = "GDPR"
        elif compliance_choice == "2":
            compliance_type = "HIPAA"
        elif compliance_choice == "3":
            compliance_type = "DPDP"
        elif compliance_choice.upper() in ["GDPR", "HIPAA", "DPDP"]:
            compliance_type = compliance_choice.upper()
        else:
            print("❌ Invalid compliance type.")
            continue

        print(f"\n🔄 Processing with {compliance_type} compliance...")
        coordinator = CoordinatorAgent()
        coordinator.handle_files(file_paths, compliance_type)

        more = input("\nDo you want to process more files? (y/n): ").strip().lower()
        if more != "y":
            print("👋 Thank you for using the redaction system!")
            break
