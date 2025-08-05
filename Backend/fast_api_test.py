import os
import json
import docx
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import re
from typing import List
import fitz
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
import time
from gliner import GLiNER
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import shutil
import tempfile
from starlette.background import BackgroundTask
import uuid
import uvicorn

app = FastAPI()

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

try:
    gliner = GLiNER.from_pretrained("./model/gliner_model")
except Exception:
    gliner = None

try:
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0,
        google_api_key=GEMINI_API_KEY,
        timeout=30
    )
except Exception:
    llm = None

class RunnerAgent:
    def load_text(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                self.json_data = json.load(f)
                text = json.dumps(self.json_data, indent=2)
        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file type. Use PDF, TXT, JSON, or DOCX.")
        return text

    def get_file_type(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            return "pdf"
        elif file_path.endswith(".txt"):
            return "txt"
        elif file_path.endswith(".json"):
            return "json"
        elif file_path.endswith(".docx"):
            return "docx"
        else:
            raise ValueError("Unsupported file type")

    def save_redacted_text(self, redacted_text: str, original_path: str, output_path: str):
        file_type = self.get_file_type(original_path)
        if file_type == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(redacted_text)
        elif file_type == "json":
            try:
                redacted_json = json.loads(redacted_text)
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(redacted_json, f, indent=2)
            except json.JSONDecodeError:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(redacted_text)
        elif file_type == "docx":
            doc = Document()
            for paragraph in redacted_text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(output_path)
        elif file_type == "pdf":
            c = canvas.Canvas(output_path, pagesize=letter)
            lines = redacted_text.split("\n")
            width, height = letter
            y = height - 40
            for line in lines:
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:100])
                y -= 15
            c.save()

class RedactorAgent:
    def __init__(self):
        self.name_exclusions = {
            'united states', 'new york', 'los angeles', 'san francisco',
            'machine learning', 'data science', 'artificial intelligence',
            'google drive', 'microsoft office', 'adobe acrobat', 'dear sir',
            'dear madam', 'yours truly', 'best regards', 'thank you',
            'monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday'
        }

    def detect_sensitive_info(self, text: str) -> List[dict]:
        all_results = []
        if gliner is not None:
            try:
                gliner_results = self._detect_with_gliner(text)
                all_results.extend(gliner_results)
            except:
                pass
        regex_results = self._regex_fallback(text)
        all_results.extend(regex_results)
        unique_items = {}
        for item in all_results:
            key = f"{item['type']}_{item['value'].lower()}"
            unique_items[key] = item
        return list(unique_items.values())

    def _detect_with_gliner(self, text: str) -> List[dict]:
        labels = ["Person", "Organization", "Date", "Email", "Phone", "Location", "URL", "Money", "Time"]
        try:
            text_chunk = text[:5000]
            entities = gliner.predict_entities(text_chunk, labels=labels, threshold=0.4)
            results = []
            for ent in entities:
                entity_type = ent["label"].lower()
                entity_value = ent["text"].strip()
                if entity_type == "person":
                    if self._is_likely_person_name(entity_value):
                        results.append({"type": "name", "value": entity_value})
                elif entity_type == "organization":
                    results.append({"type": "organization", "value": entity_value})
                elif entity_type in ["email", "phone", "url"]:
                    results.append({"type": entity_type, "value": entity_value})
                elif entity_type == "location":
                    results.append({"type": "location", "value": entity_value})
                elif entity_type in ["date", "time"]:
                    results.append({"type": "date", "value": entity_value})
                elif entity_type == "money":
                    results.append({"type": "financial", "value": entity_value})
            return results
        except:
            return []

    def _regex_fallback(self, text: str) -> List[dict]:
        patterns = {
            "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "phone": [
                r'\b\d{3}-\d{3}-\d{4}\b',
                r'\b\(\d{3}\)\s*\d{3}-\d{4}\b',
                r'\b\d{10}\b',
                r'\b\d{3}\.\d{3}\.\d{4}\b',
                r'\+\d{1,3}[\s-]?\d{3,4}[\s-]?\d{3,4}[\s-]?\d{3,4}\b'
            ],
            "url": r'https?://[^\s<>"{}|\\^`\[\]]+',
            "ssn": r'\b\d{3}-\d{2}-\d{4}\b',
            "credit_card": r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b',
            "ip_address": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
        }
        found = []
        emails = re.findall(patterns["email"], text, re.IGNORECASE)
        for email in emails:
            found.append({"type": "email", "value": email})
        for pattern in patterns["phone"]:
            phones = re.findall(pattern, text)
            for phone in phones:
                digits_only = re.sub(r'\D', '', phone)
                if len(digits_only) >= 10:
                    found.append({"type": "phone", "value": phone})
        urls = re.findall(patterns["url"], text)
        for url in urls:
            found.append({"type": "url", "value": url})
        ssns = re.findall(patterns["ssn"], text)
        for ssn in ssns:
            found.append({"type": "ssn", "value": ssn})
        cards = re.findall(patterns["credit_card"], text)
        for card in cards:
            digits_only = re.sub(r'\D', '', card)
            if len(digits_only) == 16:
                found.append({"type": "credit_card", "value": card})
        ips = re.findall(patterns["ip_address"], text)
        for ip in ips:
            octets = ip.split('.')
            if all(0 <= int(octet) <= 255 for octet in octets):
                found.append({"type": "ip_address", "value": ip})
        return found

    def _is_likely_person_name(self, name: str) -> bool:
        name_lower = name.lower()
        if name_lower in self.name_exclusions:
            return False
        non_name_patterns = [
            r'\b(mr|mrs|ms|dr|prof|sir|madam)\b',
            r'\b(inc|llc|corp|ltd|co)\b',
            r'\b(street|avenue|road|drive|lane)\b',
            r'\b(january|february|march|april|may|june|july|august|september|october|november|december)\b'
        ]
        for pattern in non_name_patterns:
            if re.search(pattern, name_lower):
                return False
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            return False
        if any(len(word) < 2 or len(word) > 20 for word in words):
            return False
        return True

    def redact(self, text: str, sensitive_items: List[dict]) -> str:
        redacted_text = text
        sorted_items = sorted(sensitive_items, key=lambda x: len(x["value"]), reverse=True)
        for item in sorted_items:
            original_value = item["value"]
            item_type = item["type"].upper()
            redaction_tag = f"[REDACTED_{item_type}]"
            if item["type"] in ["email", "url", "ssn", "credit_card", "ip_address"]:
                pattern = re.escape(original_value)
            elif item["type"] in ["name", "person", "organization", "location"]:
                escaped_value = re.escape(original_value)
                pattern = r'\b' + escaped_value + r'\b'
            elif item["type"] == "phone":
                clean_phone = re.sub(r'[\s\-\(\)]', '', original_value)
                if len(clean_phone) >= 10:
                    pattern = re.escape(original_value)
                else:
                    continue
            else:
                pattern = r'\b' + re.escape(original_value) + r'\b'
            try:
                redacted_text = re.sub(pattern, redaction_tag, redacted_text, flags=re.IGNORECASE)
            except:
                redacted_text = redacted_text.replace(original_value, redaction_tag)
        return redacted_text

    def redact_json(self, data: dict, sensitive_items: List[dict]) -> dict:
        redacted_data = json.dumps(data, indent=2)
        redacted_data = self.redact(redacted_data, sensitive_items)
        try:
            return json.loads(redacted_data)
        except json.JSONDecodeError:
            return {"redacted_content": redacted_data}

    def redact_pdf_pymupdf(self, file_path: str, sensitive_items: List[dict], output_path: str):
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                for item in sensitive_items:
                    text_instances = page.search_for(item["value"])
                    for inst in text_instances:
                        page.add_redact_annot(inst, fill=(0, 0, 0))
                page.apply_redactions()
            doc.save(output_path)
            doc.close()
        except:
            try:
                doc.close()
            except:
                pass

class PIIAgent:
    def detect_pii(self, text: str) -> list:
        patterns = {
            "email": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            "phone": r'\b(?:\+1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b',
            "ssn": r'\b\d{3}-\d{2}-\d{4}\b',
            "credit_card": r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b',
        }
        matches = []
        for name, pattern in patterns.items():
            for match in re.findall(pattern, text):
                if isinstance(match, tuple):
                    match = ''.join(match)
                matches.append({"type": name, "value": match})
        return matches

class ComplianceAgent:
    def apply_policy(self, pii_items: list, compliance_type: str) -> list:
        redactions = []
        for item in pii_items:
            if compliance_type == "GDPR":
                redactions.append(item["value"])
            elif compliance_type == "HIPAA":
                if item["type"] in ["ssn", "phone", "name", "email"]:
                    redactions.append(item["value"])
            elif compliance_type == "DPDP":
                if item["type"] != "url":
                    redactions.append(item["value"])
        return redactions

    def validate_redaction(self, redacted_text: str, compliance_type: str) -> str:
        if llm is None:
            return f"Compliance validation skipped - LLM not available. Processed with {compliance_type} standards."
        try:
            prompt = (
                f"You are a compliance officer validating text redactions for {compliance_type}. "
                f"Analyze the redacted text and check for:\n"
                f"1. Any remaining personal identifiable information\n"
                f"2. Compliance with {compliance_type} standards\n"
                f"3. Proper redaction formatting\n\n"
                f"Return a brief assessment (2-3 sentences) of compliance status.\n\n"
                f"Redacted Text:\n{redacted_text[:1000]}"
            )
            result = llm.invoke(prompt)
            return result.content if hasattr(result, 'content') else str(result)
        except Exception as e:
            return f"Compliance validation completed with basic standards. Error: {str(e)}"

class AuditAgent:
    def log_metadata(self, original_text: str, redacted_text: str, sensitive_items: List[dict], compliance_feedback: str, file_path: str) -> str:
        item_counts = {}
        for item in sensitive_items:
            item_type = item['type']
            item_counts[item_type] = item_counts.get(item_type, 0) + 1
        metadata = {
            "timestamp": time.strftime('%Y-%m-%d %H:%M:%S'),
            "file_path": file_path,
            "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            "original_length": len(original_text),
            "redacted_length": len(redacted_text),
            "total_redacted_items": len(sensitive_items),
            "redacted_items_by_type": item_counts,
            "redacted_items": [
                {"type": item["type"], "value_length": len(item["value"])}
                for item in sensitive_items
            ],
            "compliance_notes": str(compliance_feedback),
            "processing_status": "completed"
        }
        os.makedirs("audit_logs", exist_ok=True)
        log_file = f"audit_logs/audit_log_{os.path.basename(file_path)}_{int(time.time())}.json"
        with open(log_file, "w", encoding="utf-8") as f:
            json.dump(metadata, f, indent=2)
        return log_file

class CoordinatorAgent:
    def __init__(self):
        self.runner = RunnerAgent()
        self.redactor = RedactorAgent()
        self.compliance = ComplianceAgent()
        self.audit = AuditAgent()

    async def handle_file(self, file: UploadFile, compliance_type: str, temp_dir: str) -> tuple:
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in [".pdf", ".txt", ".json", ".docx"]:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, TXT, JSON, or DOCX.")
        
        temp_file_path = os.path.join(temp_dir, f"{uuid.uuid4()}{file_ext}")
        with open(temp_file_path, "wb") as f:
            f.write(await file.read())
        
        try:
            original_text = self.runner.load_text(temp_file_path)
            pii_items = self.redactor.detect_sensitive_info(original_text)
            if not pii_items:
                return None, None, file.filename
            
            redacted_text = self.redactor.redact(original_text, pii_items)
            output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_redacted{file_ext}")
            if file_ext == ".pdf":
                self.redactor.redact_pdf_pymupdf(temp_file_path, pii_items, output_path)
            else:
                self.runner.save_redacted_text(redacted_text, temp_file_path, output_path)
            
            feedback = self.compliance.validate_redaction(redacted_text, compliance_type)
            audit_path = self.audit.log_metadata(original_text, redacted_text, pii_items, feedback, temp_file_path)
            return output_path, audit_path, file.filename
        finally:
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)

class SingleFileRequest(BaseModel):
    complianceNum: str

class MultipleFileRequest(BaseModel):
    complianceNum: str

def cleanup_files(*file_paths):
    for file_path in file_paths:
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass

@app.post("/redact/single")
async def redact_single_file(file: UploadFile = File(...), request: SingleFileRequest = None):
    print(f"Received request: {request}")
    compliance_map = {"1": "GDPR", "2": "HIPAA", "3": "DPDP"}
    compliance_type = compliance_map.get(request.complianceNum)
    if not compliance_type:
        raise HTTPException(status_code=400, detail="Invalid compliance number. Use 1 (GDPR), 2 (HIPAA), or 3 (DPDP).")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        coordinator = CoordinatorAgent()
        output_path, audit_path, _ = await coordinator.handle_file(file, compliance_type, temp_dir)
        
        if not output_path:
            return {"message": "No sensitive information detected", "audit_log": None}
        
        return FileResponse(
            path=output_path,
            filename=os.path.basename(output_path),
            background=BackgroundTask(cleanup_files, output_path, audit_path),
            headers={"X-Audit-Log": audit_path}
        )

@app.post("/redact/multiple")
async def redact_multiple_files(files: List[UploadFile] = File(...), request: MultipleFileRequest = None):
    compliance_map = {"1": "GDPR", "2": "HIPAA", "3": "DPDP"}
    compliance_type = compliance_map.get(request.complianceNum)
    if not compliance_type:
        raise HTTPException(status_code=400, detail="Invalid compliance number. Use 1 (GDPR), 2 (HIPAA), or 3 (DPDP).")
    
    with tempfile.TemporaryDirectory() as temp_dir:
        coordinator = CoordinatorAgent()
        results = []
        files_to_cleanup = []
        
        for file in files:
            output_path, audit_path, filename = await coordinator.handle_file(file, compliance_type, temp_dir)
            if output_path:
                results.append({
                    "filename": filename,
                    "redacted_file": os.path.basename(output_path),
                    "audit_log": os.path.basename(audit_path)
                })
                files_to_cleanup.extend([output_path, audit_path])
        
        if not results:
            return {"message": "No sensitive information detected in any files", "files": []}
        
        # Create a zip file containing all redacted files and audit logs
        zip_path = os.path.join(temp_dir, f"redacted_files_{int(time.time())}.zip")
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for result in results:
                zipf.write(result["redacted_file"], f"redacted/{result['redacted_file']}")
                zipf.write(result["audit_log"], f"audit_logs/{result['audit_log']}")
        
        return FileResponse(
            path=zip_path,
            filename=os.path.basename(zip_path),
            background=BackgroundTask(cleanup_files, zip_path, *files_to_cleanup),
            headers={"Content-Disposition": f"attachment; filename={os.path.basename(zip_path)}"}
        )
    
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)