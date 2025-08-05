from flask import current_app
import os
import json
import docx
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import re
from typing import List

class RunnerAgent:
    def load_text(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())

        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif file_path.endswith(".json"):
            with open(file_path, "r", encoding="utf-8") as f:
                self.json_data = json.load(f)
                text = json.dumps(self.json_data, indent=2)

        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        else:
            raise ValueError("Unsupported file type. Use PDF, TXT, JSON, or DOCX.")

        return text

    def get_file_type(self, file_path: str) -> str:
        if file_path.endswith(".pdf"):
            return "pdf"
        elif file_path.endswith(".txt"):
            return "txt"
        elif file_path.endswith(".json"):
            return "json"
        elif file_path.endswith(".docx"):
            return "docx"
        else:
            raise ValueError("Unsupported file type")

    def save_redacted_text(self, redacted_text: str, original_path: str, output_path: str):
        file_type = self.get_file_type(original_path)

        if file_type == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(redacted_text)

        elif file_type == "json":
            try:
                redacted_json = json.loads(redacted_text)
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(redacted_json, f, indent=2)
            except json.JSONDecodeError:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(redacted_text)

        elif file_type == "docx":
            doc = Document()
            for paragraph in redacted_text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(output_path)

        elif file_type == "pdf":
            c = canvas.Canvas(output_path, pagesize=letter)
            lines = redacted_text.split("\n")
            width, height = letter
            y = height - 40
            for line in lines:
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, line[:100])  # Limit line length
                y -= 15
            c.save()
